import logging
import pdfplumber
from docx import Document
from io import BytesIO

log = logging.getLogger("curator.resume_parser")

MAX_FILE_MB = 10
MAX_FILE_BYTES = MAX_FILE_MB * 1024 * 1024


def _extract_pdf_text(file_bytes: bytes) -> str:
    """Try pdfplumber first, fall back to pypdf."""
    # Method 1: pdfplumber
    try:
        with pdfplumber.open(BytesIO(file_bytes)) as pdf:
            pages = [page.extract_text() or "" for page in pdf.pages]
        text = "\n".join(pages).strip()
        if text:
            log.info("pdfplumber extracted %d chars", len(text))
            return text
        log.warning("pdfplumber returned empty text, trying pypdf fallback")
    except Exception as e:
        log.warning("pdfplumber failed: %s — trying pypdf fallback", e)

    # Method 2: pypdf fallback
    try:
        from pypdf import PdfReader
        reader = PdfReader(BytesIO(file_bytes))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            log.info("pypdf extracted %d chars", len(text))
            return text
        log.warning("pypdf also returned empty text")
    except ImportError:
        log.warning("pypdf not installed, skipping fallback")
    except Exception as e:
        log.warning("pypdf failed: %s", e)

    return ""


def parse_resume(file_bytes: bytes, filename: str) -> str:
    """
    Convert PDF, DOCX, or TXT bytes to plain text.
    Raises ValueError with a user-friendly message on failure.
    """
    if len(file_bytes) > MAX_FILE_BYTES:
        raise ValueError(f"File too large (max {MAX_FILE_MB} MB).")

    name = filename.lower().strip()
    log.info("parse_resume: %s (%d bytes)", filename, len(file_bytes))

    # ── PDF ──────────────────────────────────────────────────────────────────
    if name.endswith(".pdf"):
        text = _extract_pdf_text(file_bytes)
        if not text:
            raise ValueError(
                "Could not extract text from this PDF.\n"
                "Common causes:\n"
                "• Scanned/image-based PDF — try converting to text-based PDF\n"
                "• Password-protected PDF — remove the password first\n"
                "• Try saving as DOCX from Word or Google Docs instead"
            )
        return text

    # ── DOCX ─────────────────────────────────────────────────────────────────
    if name.endswith(".docx"):
        try:
            doc = Document(BytesIO(file_bytes))
        except Exception as e:
            raise ValueError(f"Could not open DOCX: {e}")

        parts = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        parts.append(cell.text.strip())

        text = "\n".join(parts).strip()
        if not text:
            raise ValueError("No text found in DOCX. Try re-saving the file.")
        log.info("parse_resume: DOCX extracted %d chars", len(text))
        return text

    # ── TXT ──────────────────────────────────────────────────────────────────
    if name.endswith(".txt"):
        try:
            text = file_bytes.decode("utf-8", errors="ignore").strip()
        except Exception as e:
            raise ValueError(f"Could not read text file: {e}")
        if not text:
            raise ValueError("Text file is empty.")
        log.info("parse_resume: TXT extracted %d chars", len(text))
        return text

    raise ValueError(
        f"Unsupported file type. Please upload PDF, DOCX, or TXT."
    )
